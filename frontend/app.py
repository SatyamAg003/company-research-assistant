import streamlit as st
import requests
import os
import base64
import json
import speech_recognition as sr
import pyttsx3
import threading
import queue
import tempfile
from io import BytesIO
import re

BACKEND_URL = os.getenv('BACKEND_URL', 'http://localhost:8000')

st.set_page_config(
    page_title='Company Research Assistant',
    page_icon='💼',
    layout='centered'
)

st.title('💼 Company Research Assistant')

# Initialize session state
if 'messages' not in st.session_state:
    st.session_state.messages = [
        {
            "role": "assistant",
            "content": "Hi! I'm your Company Research Assistant. I can help you research any company and generate editable account plans. Just tell me which company you'd like to research!"
        }
    ]

if 'research_in_progress' not in st.session_state:
    st.session_state.research_in_progress = False

if 'current_company' not in st.session_state:
    st.session_state.current_company = None

if 'account_plan' not in st.session_state:
    st.session_state.account_plan = None

if 'show_account_plan' not in st.session_state:
    st.session_state.show_account_plan = False

if 'research_data' not in st.session_state:
    st.session_state.research_data = None

if 'listening' not in st.session_state:
    st.session_state.listening = False

if 'voice_input' not in st.session_state:
    st.session_state.voice_input = ""

# Initialize text-to-speech engine
@st.cache_resource
def init_tts():
    try:
        engine = pyttsx3.init()
        # Configure voice settings
        voices = engine.getProperty('voices')
        if voices:
            engine.setProperty('voice', voices[0].id)  # Use first available voice
        engine.setProperty('rate', 150)  # Speed of speech
        engine.setProperty('volume', 0.8)  # Volume level
        return engine
    except:
        return None

tts_engine = init_tts()

def speak_text(text):
    """Convert text to speech in a separate thread"""
    if tts_engine:
        def speak():
            try:
                tts_engine.say(text)
                tts_engine.runAndWait()
            except Exception as e:
                print(f"TTS error: {e}")
        
        thread = threading.Thread(target=speak)
        thread.daemon = True
        thread.start()

def transcribe_audio():
    """Transcribe speech to text using microphone"""
    try:
        recognizer = sr.Recognizer()
        microphone = sr.Microphone()
        
        with microphone as source:
            st.info("🎤 Listening... Speak now!")
            recognizer.adjust_for_ambient_noise(source)
            audio = recognizer.listen(source, timeout=10, phrase_time_limit=15)
        
        try:
            text = recognizer.recognize_google(audio)
            return text
        except sr.UnknownValueError:
            return "Sorry, I couldn't understand the audio."
        except sr.RequestError:
            return "Sorry, there was an error with the speech recognition service."
            
    except Exception as e:
        return f"Error with microphone: {str(e)}"

def contains_gibberish(text):
    """Check if text contains obvious gibberish patterns"""
    text_lower = text.lower()
    
    # Common gibberish patterns (keyboard mashing, lorem ipsum, etc.)
    gibberish_patterns = [
        'asdf', 'jkl', 'qwerty', 'zxcv', 'lorem', 'ipsum', 'dolor', 'sit', 'amet',
        'test', 'example', 'sample', 'random', 'foo', 'bar', 'baz'
    ]
    
    # Check for repeating characters (like "aaaa", "jjjj", etc.)
    if re.search(r'(.)\1{3,}', text_lower):  # 4 or more repeating characters
        return True
    
    # Check for keyboard row patterns (qwerty, asdf, etc.)
    if any(pattern in text_lower for pattern in gibberish_patterns):
        return True
    
    # Check for unusual character sequences (too many consonants/vowels in a row)
    words = text_lower.split()
    for word in words:
        if len(word) > 3:
            # Check for consonant-heavy or vowel-heavy sequences
            consonants = len(re.findall(r'[bcdfghjklmnpqrstvwxyz]', word))
            vowels = len(re.findall(r'[aeiou]', word))
            if consonants > vowels * 2 or vowels > consonants * 2:  # Unbalanced ratio
                return True
    
    # Check for meaningless short words that aren't real words
    common_words = ['the', 'and', 'for', 'are', 'but', 'not', 'you', 'all', 'can', 'had', 'her', 'was', 'one', 'our', 'out', 'get', 'has', 'him', 'his', 'how', 'man', 'new', 'now', 'old', 'see', 'two', 'way', 'who', 'boy', 'did', 'its', 'let', 'put', 'say', 'she', 'too', 'use']
    if len(text_lower) < 4 and text_lower not in common_words:
        return True
    
    return False

def is_valid_company_name(company_name):
    """Validate if the extracted company name is likely to be a real company"""
    company_lower = company_name.lower()
    
    # Check for obvious gibberish
    if contains_gibberish(company_name):
        return False
    
    # Check for meaningless phrases
    meaningless_phrases = [
        'from my dreams', 'from my imagination', 'my dreams', 'my imagination',
        'something about', 'some company', 'any company', 'what were', 'were doing',
        'cousin works', 'says coffee', 'coffee great'
    ]
    
    if any(phrase in company_lower for phrase in meaningless_phrases):
        return False
    
    # Check if it's too short or too long to be reasonable
    if len(company_name) < 2 or len(company_name) > 50:
        return False
    
    # Be more lenient with capitalization for single-word company names
    words = company_name.split()
    if len(words) == 1:
        # Single word companies can be any capitalization
        return True
    else:
        # For multi-word, check if most words start with capital letters
        capitalized_words = sum(1 for word in words if word and word[0].isupper())
        if capitalized_words < len(words) * 0.3:  # Reduced from 0.5 to 0.3
            return False
    
    return True
def is_research_request(prompt):
    """Check if the user is requesting new research"""
    research_keywords = [
        'research', 'analyze', 'study', 'look up', 'find info', 
        'tell me about', 'generate account plan', 'create account plan',
        'make account plan', 'build account plan'
    ]
    
    prompt_lower = prompt.lower()
    
    # Check for explicit research keywords
    if any(keyword in prompt_lower for keyword in research_keywords):
        return True
    
    # Check if it's asking about a new company we haven't researched
    if st.session_state.current_company:
        words = prompt_lower.split()
        for word in words:
            if word not in research_keywords and len(word) > 2:
                if word != st.session_state.current_company.lower() and word.title() != st.session_state.current_company:
                    if any(ctx in prompt_lower for ctx in ['what about', 'how about', 'can you', 'what if']):
                        return True
    
    return False

def extract_company_name(prompt):
    """Extract company name from research request"""
    research_keywords = [
        'research', 'analyze', 'study', 'look up', 'find info', 
        'tell me about', 'generate account plan', 'create account plan',
        'make account plan', 'build account plan', 'account plan for',
        'research on', 'analyze on'
    ]
    
    prompt_lower = prompt.lower()
    
    # FIRST: Try to extract company name even from messy input
    # Look for research patterns followed by company names
    for keyword in research_keywords:
        if keyword in prompt_lower:
            # Find the position after the research keyword
            keyword_pos = prompt_lower.find(keyword)
            after_keyword = prompt_lower[keyword_pos + len(keyword):].strip()
            
            # Extract the first meaningful word(s) after research keyword
            words_after = after_keyword.split()
            potential_company_words = []
            
            for word in words_after:
                if (word not in ['the', 'a', 'an', 'about', 'on', 'for', 'and', 'but', 'anyway'] and
                    len(word) > 2 and not contains_gibberish(word)):
                    potential_company_words.append(word)
                    # Stop if we hit conversational markers
                    if word in ['my', 'i', 'we', 'you', 'anyway', 'what', 'were']:
                        break
            
            if potential_company_words:
                company_name = ' '.join(potential_company_words).title()
                if is_valid_company_name(company_name):
                    return company_name
    
    # SECOND: Check for vague requests that need clarification
    vague_phrases = [
        'something about',
        'some companies',
        'any company', 
        'a company',
        'companies in general',
        'business in general',
        'tell me about companies',
        'research companies'
    ]
    
    # If it's a vague request without any specific company mention, return None
    if any(phrase in prompt_lower for phrase in vague_phrases):
        # But only if no company is mentioned
        company_indicators = ['microsoft', 'apple', 'google', 'amazon', 'tesla', 'netflix', 
                             'meta', 'facebook', 'ibm', 'intel', 'samsung', 'sony']
        if not any(company in prompt_lower for company in company_indicators):
            return None
    
    # THIRD: Fallback - extract company using the old method but be more lenient
    words = prompt_lower.split()
    
    # Remove research keywords and common words
    filtered_words = []
    for word in words:
        if (word not in research_keywords and 
            word not in ['the', 'a', 'an', 'about', 'on', 'for', 'something', 'companies', 'business',
                        'lorem', 'ipsum', 'from', 'my', 'dreams', 'corporation', 'company',
                        'cousin', 'works', 'says', 'coffee', 'great', 'anyway', 'what', 'were', 'doing'] and
            len(word) > 1 and  # Ignore single letters
            not word.isnumeric() and  # Ignore numbers
            not contains_gibberish(word)):  # Ignore gibberish
            filtered_words.append(word)
    
    company_name = ' '.join(filtered_words).title()
    
    # Additional validation
    if len(company_name.strip()) < 2:
        return None
    
    # Final validation check
    if not is_valid_company_name(company_name):
        return None
    
    return company_name.strip()

def is_account_plan_request(prompt):
    """Check if user specifically wants an account plan"""
    plan_keywords = [
        'account plan', 'generate plan', 'create plan', 'make plan',
        'build plan', 'business plan', 'strategy plan'
    ]
    return any(keyword in prompt.lower() for keyword in plan_keywords)

# Voice input handling
if st.session_state.listening:
    with st.spinner("🎤 Listening... Speak now!"):
        try:
            transcribed_text = transcribe_audio()
            if transcribed_text and not transcribed_text.startswith("Sorry") and not transcribed_text.startswith("Error"):
                st.session_state.voice_input = transcribed_text
                st.session_state.listening = False
                st.rerun()
            else:
                st.error(transcribed_text)
                st.session_state.listening = False
        except Exception as e:
            st.error(f"Voice input failed: {str(e)}")
            st.session_state.listening = False

# Display chat messages
for message in st.session_state.messages:
    with st.chat_message(message["role"]):
        st.markdown(message["content"])
        
        # Add text-to-speech button for assistant messages
        if message["role"] == "assistant" and tts_engine:
            if st.button("🔊 Speak", key=f"speak_{hash(message['content'])}"):
                speak_text(message["content"])

# Voice input section
col1, col2 = st.columns([3, 1])

with col1:
    # Text input
    prompt = st.chat_input("Ask me to research a company or ask follow-up questions...")

with col2:
    # Voice input button
    if st.button("🎤 Voice Input", use_container_width=True):
        st.session_state.listening = True
        st.rerun()

# Use voice input if available
if st.session_state.voice_input and not prompt:
    prompt = st.session_state.voice_input
    st.session_state.voice_input = ""  # Reset after use

# Process input
if prompt:
    # Add user message to chat history
    st.session_state.messages.append({"role": "user", "content": prompt})
    with st.chat_message("user"):
        st.markdown(prompt)

    # Determine the type of request
    should_generate_account_plan = is_account_plan_request(prompt)
    should_do_research = is_research_request(prompt)
    
    company_to_research = None
    if should_do_research:
        company_to_research = extract_company_name(prompt)

    # Generate response
    with st.chat_message("assistant"):
        message_placeholder = st.empty()
        
        # Handle vague research requests
        if should_do_research and not company_to_research:
            clarification_response = """
I'd be happy to help you research companies! To get started, I need to know which specific company you're interested in.

**Examples:**
- "Research Microsoft"
- "Tell me about Tesla"
- "Analyze Apple's business strategy"
- "Generate account plan for Amazon"

Which specific company would you like me to research?
"""
            message_placeholder.markdown(clarification_response)
            st.session_state.messages.append({"role": "assistant", "content": clarification_response})
            
        elif should_do_research and company_to_research:
            # Start research
            st.session_state.research_in_progress = True
            st.session_state.current_company = company_to_research
            
            message_placeholder.markdown(f"🔍 Starting research on **{company_to_research}**...")
            
            try:
                # Call research endpoint
                research_response = requests.post(
                    f"{BACKEND_URL}/api/research",
                    json={"company": company_to_research, "fetch_news": True},
                    timeout=120
                )
                
                if research_response.status_code == 200:
                    research_data = research_response.json()
                    st.session_state.research_data = research_data["data"]
                    
                    # Show research updates
                    updates = research_data.get("updates", [])
                    research_updates_text = "\n".join([f"• {update}" for update in updates])
                    
                    # Generate account plan only if specifically requested
                    if should_generate_account_plan:
                        message_placeholder.markdown(f"🔍 Research completed! 📊 Generating account plan...")
                        
                        account_plan_response = requests.post(
                            f"{BACKEND_URL}/api/generate-account-plan",
                            json={
                                "company": company_to_research,
                                "research_data": research_data["data"]
                            }
                        )
                        
                        if account_plan_response.status_code == 200:
                            account_plan_data = account_plan_response.json()
                            
                            if "error" not in account_plan_data:
                                st.session_state.account_plan = account_plan_data
                                st.session_state.show_account_plan = True
                                
                                final_response = f"""
## 📊 Research Complete: {company_to_research}

✅ I've generated a comprehensive account plan based on research from Wikipedia, DuckDuckGo, and GNews.

**🔍 Research Process:**
{research_updates_text}

👇 **You can now review and edit the account plan below!** Each section is editable - just click on any section to modify it.
"""
                            else:
                                final_response = f"## 📊 Research Complete: {company_to_research}\n\n{research_updates_text}\n\n⚠️ Could not generate account plan: {account_plan_data['error']}"
                        else:
                            final_response = f"## 📊 Research Complete: {company_to_research}\n\n{research_updates_text}\n\n⚠️ Account plan generation failed."
                    else:
                        # Just research, no account plan
                        final_response = f"""
## 🔍 Research Complete: {company_to_research}

✅ I've gathered comprehensive information about {company_to_research} from Wikipedia, DuckDuckGo, and GNews.

**Research Process:**
{research_updates_text}

You can now ask me specific questions about {company_to_research}, or say "generate account plan" to create a detailed business plan.
"""
                    
                    message_placeholder.markdown(final_response)
                    st.session_state.messages.append({"role": "assistant", "content": final_response})
                    
                else:
                    error_msg = f"❌ Sorry, I couldn't research {company_to_research}. Please try again later."
                    message_placeholder.markdown(error_msg)
                    st.session_state.messages.append({"role": "assistant", "content": error_msg})
                    
            except Exception as e:
                error_msg = f"❌ Research error: {str(e)}"
                message_placeholder.markdown(error_msg)
                st.session_state.messages.append({"role": "assistant", "content": error_msg})
            
            st.session_state.research_in_progress = False
            
        elif should_generate_account_plan and st.session_state.research_data and st.session_state.current_company:
            # Generate account plan for already researched company
            message_placeholder.markdown(f"📊 Generating account plan for {st.session_state.current_company}...")
            
            try:
                account_plan_response = requests.post(
                    f"{BACKEND_URL}/api/generate-account-plan",
                    json={
                        "company": st.session_state.current_company,
                        "research_data": st.session_state.research_data
                    }
                )
                
                if account_plan_response.status_code == 200:
                    account_plan_data = account_plan_response.json()
                    
                    if "error" not in account_plan_data:
                        st.session_state.account_plan = account_plan_data
                        st.session_state.show_account_plan = True
                        
                        final_response = f"""
## 📊 Account Plan Generated: {st.session_state.current_company}

✅ I've created a comprehensive account plan based on our previous research.

👇 **You can now review and edit the account plan below!** Each section is editable.
"""
                    else:
                        final_response = f"❌ Could not generate account plan: {account_plan_data['error']}"
                else:
                    final_response = "❌ Account plan generation failed."
                
                message_placeholder.markdown(final_response)
                st.session_state.messages.append({"role": "assistant", "content": final_response})
                
            except Exception as e:
                error_msg = f"❌ Error generating account plan: {str(e)}"
                message_placeholder.markdown(error_msg)
                st.session_state.messages.append({"role": "assistant", "content": error_msg})
            
        else:
            # Regular chat message - no research, no account plan generation
            message_placeholder.markdown("🤔 Thinking...")
            
            try:
                # Prepare research data for chat if we have it
                research_data_for_chat = None
                if st.session_state.research_data and st.session_state.current_company:
                    research_data_for_chat = st.session_state.research_data
                    research_data_for_chat["company"] = st.session_state.current_company
                
                chat_response = requests.post(
                    f"{BACKEND_URL}/api/chat",
                    json={
                        "message": prompt,
                        "conversation_history": st.session_state.messages
                    },
                    timeout=60
                )
                
                if chat_response.status_code == 200:
                    chat_data = chat_response.json()
                    response_text = chat_data["response"]
                    
                    # Add research context indicator if available
                    if chat_data["research_available"] and st.session_state.current_company:
                        response_text = f"*[Using research data for {st.session_state.current_company}]*\n\n{response_text}"
                    
                    message_placeholder.markdown(response_text)
                    st.session_state.messages.append({"role": "assistant", "content": response_text})
                    
                    # Auto-speak the response if it's not too long
                    if len(response_text) < 500:  # Don't speak very long responses
                        speak_text(response_text)
                        
                else:
                    error_msg = "❌ Sorry, I'm having trouble responding right now. Please try again."
                    message_placeholder.markdown(error_msg)
                    st.session_state.messages.append({"role": "assistant", "content": error_msg})
                    
            except Exception as e:
                error_msg = f"❌ Connection error: {str(e)}"
                message_placeholder.markdown(error_msg)
                st.session_state.messages.append({"role": "assistant", "content": error_msg})

# Account Plan Editor (only shown when specifically requested and generated)
if st.session_state.account_plan and st.session_state.show_account_plan:
    st.markdown("---")
    st.header(f"📋 Account Plan: {st.session_state.current_company}")
    st.info("💡 You can edit any section below. Your changes will be saved automatically.")
    
    # Editable sections
    col1, col2 = st.columns(2)
    
    with col1:
        st.subheader("Executive Summary")
        st.session_state.account_plan['executive_summary'] = st.text_area(
            "Executive Summary",
            value=st.session_state.account_plan['executive_summary'],
            height=150,
            key="exec_summary_edit"
        )
        
        st.subheader("Company Overview")
        st.session_state.account_plan['company_overview'] = st.text_area(
            "Company Overview",
            value=st.session_state.account_plan['company_overview'],
            height=150,
            key="company_overview_edit"
        )
        
        st.subheader("Key Contacts")
        st.session_state.account_plan['key_contacts'] = st.text_area(
            "Key Contacts",
            value=st.session_state.account_plan['key_contacts'],
            height=120,
            key="key_contacts_edit"
        )
    
    with col2:
        st.subheader("Strengths & Weaknesses")
        st.session_state.account_plan['strengths_weaknesses'] = st.text_area(
            "Strengths & Weaknesses",
            value=st.session_state.account_plan['strengths_weaknesses'],
            height=150,
            key="strengths_weaknesses_edit"
        )
        
        st.subheader("Opportunities & Risks")
        st.session_state.account_plan['opportunities_risks'] = st.text_area(
            "Opportunities & Risks",
            value=st.session_state.account_plan['opportunities_risks'],
            height=150,
            key="opportunities_risks_edit"
        )
        
        st.subheader("Engagement Plan")
        st.session_state.account_plan['engagement_plan'] = st.text_area(
            "Engagement Plan",
            value=st.session_state.account_plan['engagement_plan'],
            height=120,
            key="engagement_plan_edit"
        )
    
    # Download options
    st.markdown("---")
    st.subheader("💾 Export Options")
    
    col1, col2, col3 = st.columns(3)
    
    with col1:
        if st.button("📥 Download as DOCX"):
            try:
                from docx import Document
                doc = Document()
                doc.add_heading(f'Account Plan: {st.session_state.current_company}', level=1)
                
                # Add all sections
                sections_display = {
                    'executive_summary': 'Executive Summary',
                    'company_overview': 'Company Overview', 
                    'key_contacts': 'Key Contacts',
                    'strengths_weaknesses': 'Strengths & Weaknesses',
                    'opportunities_risks': 'Opportunities & Risks',
                    'engagement_plan': 'Engagement Plan'
                }
                
                for key, display_name in sections_display.items():
                    doc.add_heading(display_name, level=2)
                    doc.add_paragraph(st.session_state.account_plan[key])
                
                filename = f"account_plan_{st.session_state.current_company.replace(' ', '_').lower()}.docx"
                doc.save(filename)
                
                with open(filename, 'rb') as f:
                    data = f.read()
                b64 = base64.b64encode(data).decode()
                href = f'<a href="data:application/octet-stream;base64,{b64}" download="{filename}">📥 Download DOCX</a>'
                st.markdown(href, unsafe_allow_html=True)
                st.success('Document saved successfully!')
                
            except ImportError:
                st.error('python-docx not installed. Run: pip install python-docx')
            except Exception as e:
                st.error(f'Failed to create document: {str(e)}')
    
    with col2:
        if st.button("📋 Copy to Clipboard"):
            try:
                plan_text = f"ACCOUNT PLAN: {st.session_state.current_company}\n\n"
                for section, content in st.session_state.account_plan.items():
                    plan_text += f"{section.replace('_', ' ').title()}:\n{content}\n\n"
                
                st.code(plan_text)
                st.success("Plan copied to clipboard! (Use Ctrl+C to copy)")
            except Exception as e:
                st.error(f"Failed to copy: {str(e)}")
    
    with col3:
        if st.button("❌ Hide Plan"):
            st.session_state.show_account_plan = False
            st.rerun()

# Sidebar with information
with st.sidebar:
    st.header("🎤 Voice Features")
    
    # Voice settings
    st.subheader("Voice Settings")
    auto_speak = st.checkbox("Auto-speak responses", value=True)
    voice_rate = st.slider("Speech Rate", 100, 200, 150)
    
    if tts_engine:
        tts_engine.setProperty('rate', voice_rate)
    
    st.markdown("---")
    st.header("💡 How to Use")
    st.markdown("""
    **Voice Commands:**
    - Click 🎤 button and speak naturally
    - "Research Microsoft"
    - "Generate account plan"
    - "Tell me about their products"
    
    **Text Input:**
    - Type your questions normally
    - Use 🔊 to hear responses
    
    **Research Sources:**
    - 📚 Wikipedia
    - 🌐 DuckDuckGo  
    - 📰 GNews (Recent news)
    """)
    
    st.markdown("---")
    
    if st.session_state.current_company:
        st.success(f"✅ Researched: **{st.session_state.current_company}**")
        
        if st.session_state.research_data and not st.session_state.show_account_plan:
            if st.button("📊 Generate Account Plan"):
                st.session_state.messages.append({
                    "role": "user", 
                    "content": "generate account plan"
                })
                st.rerun()
    
    if st.button("🗑️ Clear Conversation"):
        st.session_state.messages = [
            {
                "role": "assistant",
                "content": "Hi! I'm your Company Research Assistant. I can help you research any company and generate editable account plans. Just tell me which company you'd like to research!"
            }
        ]
        st.session_state.current_company = None
        st.session_state.research_in_progress = False
        st.session_state.account_plan = None
        st.session_state.show_account_plan = False
        st.session_state.research_data = None
        st.session_state.listening = False
        st.session_state.voice_input = ""
        st.rerun()

    # Installation instructions
    with st.expander("🔧 Voice Setup Requirements"):
        st.markdown("""
        **For Voice Input:**
        ```bash
        pip install SpeechRecognition pyttsx3 pyaudio
        ```
        
        **On Windows:**
        - PyAudio should install automatically
        
        **On Mac:**
        ```bash
        brew install portaudio
        pip install pyaudio
        ```
        
        **On Linux:**
        ```bash
        sudo apt-get install python3-pyaudio
        pip install pyaudio
        ```
        """)
